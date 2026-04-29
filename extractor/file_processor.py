# extractor/file_processor.py
# ═══════════════════════════════════════════════
# Handles reading and converting different
# file types into plain text for AI extraction
# Supports: PDF, Excel, CSV, Word, Text
# ═══════════════════════════════════════════════

import io
import re
from pathlib import Path
from config import SUPPORTED_EXTENSIONS


class FileProcessor:

    def __init__(self):
        self.stats = {
            "pdf":     0,
            "excel":   0,
            "csv":     0,
            "word":    0,
            "text":    0,
            "skipped": 0,
        }

    def get_file_type(self, filename):
        """Return file type string from extension"""
        ext = Path(filename).suffix.lower()
        return SUPPORTED_EXTENSIONS.get(ext, None)

    def extract_text(self, file_bytes, filename):
        """
        Main entry point.
        Detects file type and extracts text.
        Returns plain text string.
        """
        file_type = self.get_file_type(filename)

        if file_type is None:
            self.stats["skipped"] += 1
            return None, "unsupported"

        if file_type == "pdf":
            return self._process_pdf(file_bytes, filename)

        elif file_type == "excel":
            return self._process_excel(file_bytes, filename)

        elif file_type == "csv":
            return self._process_csv(file_bytes, filename)

        elif file_type == "word":
            return self._process_word(file_bytes, filename)

        elif file_type == "text":
            return self._process_text(file_bytes, filename)

        return None, "unsupported"

    # ══════════════════════════════════════
    #  PDF PROCESSING
    # ══════════════════════════════════════
    def _process_pdf(self, file_bytes, filename):
        """
        Try multiple PDF extraction methods.
        Returns (text, method_used)
        """
        self.stats["pdf"] += 1

        # Method 1 — pdfplumber (best for tables)
        try:
            import pdfplumber
            text = self._pdf_with_pdfplumber(file_bytes)
            if text and len(text.strip()) > 100:
                return text, "pdfplumber"
        except ImportError:
            pass
        except Exception as e:
            print(f"    pdfplumber failed: {e}")

        # Method 2 — PyMuPDF (fast, good for text)
        try:
            import fitz
            text = self._pdf_with_pymupdf(file_bytes)
            if text and len(text.strip()) > 100:
                return text, "pymupdf"
        except ImportError:
            pass
        except Exception as e:
            print(f"    PyMuPDF failed: {e}")

        # Method 3 — pypdf (basic fallback)
        try:
            from pypdf import PdfReader
            text = self._pdf_with_pypdf(file_bytes)
            if text and len(text.strip()) > 50:
                return text, "pypdf"
        except ImportError:
            pass
        except Exception as e:
            print(f"    pypdf failed: {e}")

        # Method 4 — Raw bytes decode (last resort)
        try:
            raw = file_bytes.decode("utf-8", errors="ignore")
            cleaned = self._clean_pdf_raw(raw)
            if cleaned and len(cleaned.strip()) > 50:
                return cleaned, "raw"
        except Exception:
            pass

        return None, "failed"

    def _pdf_with_pdfplumber(self, file_bytes):
        """Extract PDF using pdfplumber — best for tables"""
        import pdfplumber
        text_parts = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text() or ""
                if page_text:
                    text_parts.append(page_text)

                # Extract tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            clean_row = [
                                str(cell or "").strip()
                                for cell in row
                            ]
                            text_parts.append(
                                " | ".join(clean_row)
                            )

        return "\n".join(text_parts)

    def _pdf_with_pymupdf(self, file_bytes):
        """Extract PDF using PyMuPDF (fitz)"""
        import fitz
        text_parts = []

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()

        return "\n".join(text_parts)

    def _pdf_with_pypdf(self, file_bytes):
        """Extract PDF using pypdf"""
        from pypdf import PdfReader
        text_parts = []

        reader = PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n".join(text_parts)

    def _clean_pdf_raw(self, raw_text):
        """Clean raw PDF byte string"""
        # Remove binary garbage, keep readable text
        cleaned = re.sub(r'[^\x20-\x7E\n\t]', ' ', raw_text)
        cleaned = re.sub(r'\s{4,}', '  ', cleaned)
        return cleaned

    # ══════════════════════════════════════
    #  EXCEL PROCESSING
    # ══════════════════════════════════════
    def _process_excel(self, file_bytes, filename):
        """Extract Excel data as structured text"""
        import pandas as pd
        self.stats["excel"] += 1

        try:
            xl        = pd.ExcelFile(io.BytesIO(file_bytes))
            text_parts = []

            for sheet_name in xl.sheet_names[:8]:
                try:
                    df = pd.read_excel(
                        io.BytesIO(file_bytes),
                        sheet_name=sheet_name,
                        header=None,
                    )

                    text_parts.append(f"\n=== Sheet: {sheet_name} ===")

                    # Find rows with numeric values (prices)
                    for idx, row in df.iterrows():
                        row_str = " | ".join([
                            str(v).strip()
                            for v in row
                            if str(v).strip()
                            and str(v).strip() != "nan"
                        ])
                        if row_str:
                            text_parts.append(row_str)

                except Exception as e:
                    text_parts.append(
                        f"Sheet {sheet_name}: read error {e}"
                    )

            return "\n".join(text_parts), "pandas_excel"

        except Exception as e:
            return None, f"excel_failed: {e}"

    # ══════════════════════════════════════
    #  CSV PROCESSING
    # ══════════════════════════════════════
    def _process_csv(self, file_bytes, filename):
        """Extract CSV data as text"""
        import pandas as pd
        self.stats["csv"] += 1

        try:
            # Try different encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    df = pd.read_csv(
                        io.BytesIO(file_bytes),
                        encoding=encoding,
                        on_bad_lines="skip"
                    )
                    text = df.to_string(
                        index=False, max_rows=200
                    )
                    return text, "pandas_csv"
                except UnicodeDecodeError:
                    continue

        except Exception as e:
            pass

        # Fallback — raw text
        for encoding in ["utf-8", "latin-1"]:
            try:
                text = file_bytes.decode(encoding)
                return text[:10000], "raw_csv"
            except UnicodeDecodeError:
                continue

        return None, "csv_failed"

    # ══════════════════════════════════════
    #  WORD PROCESSING
    # ══════════════════════════════════════
    def _process_word(self, file_bytes, filename):
        """Extract Word document text"""
        self.stats["word"] += 1

        try:
            from docx import Document
            doc        = Document(io.BytesIO(file_bytes))
            text_parts = []

            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ])
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts), "python_docx"

        except ImportError:
            pass
        except Exception as e:
            pass

        # Fallback
        try:
            raw = file_bytes.decode("utf-8", errors="ignore")
            return raw[:8000], "raw_word"
        except Exception:
            return None, "word_failed"

    # ══════════════════════════════════════
    #  TEXT PROCESSING
    # ══════════════════════════════════════
    def _process_text(self, file_bytes, filename):
        """Read plain text files"""
        self.stats["text"] += 1

        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                text = file_bytes.decode(encoding)
                return text[:10000], "text"
            except UnicodeDecodeError:
                continue

        return None, "text_failed"

    # ══════════════════════════════════════
    #  HELPERS
    # ══════════════════════════════════════
    def get_stats(self):
        """Return processing statistics"""
        return self.stats

    def is_supported(self, filename):
        """Check if file type is supported"""
        ext = Path(filename).suffix.lower()
        return ext in SUPPORTED_EXTENSIONS
